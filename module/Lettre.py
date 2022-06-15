import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from module.Person import Person
from docx.shared import Cm


class Letter:

    def save_letter(self, expediteur, destinataire, obj, msg):
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(3.81)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        p_expediteur = document.add_paragraph(self.check(expediteur))

        p_destinataire = document.add_paragraph(f"{self.check(destinataire,True)}\n\n")
        p_destinataire.paragraph_format.left_indent = Cm(9)
        # p_destinataire.paragraph_format.space_before = Cm(11)

        p_lieu = document.add_paragraph(
            f"{expediteur.ville},\nle {datetime.datetime.now().strftime('%d/%m/%Y')}")
        p_lieu.paragraph_format.left_indent = Cm(9)

        if obj != "":
            p_objet = document.add_paragraph(f"Objet : {obj}")

        p_appel = document.add_paragraph(f"Madame, Monsieur,\n\n")
        msg_split = msg.split("\n")
        
        for text in msg_split:
            document.add_paragraph(text)

        p_signature = document.add_paragraph(
            f"{expediteur.prenom} {expediteur.nom}")
        p_signature.paragraph_format.left_indent = Cm(9)
        try:
            document.save("letter.docx")
            return True
        except:
            return False
    
    def check(self, person,dest=False):
        text = f"{person.nom} {person.prenom}\n{person.adresse}\n{person.cp} {person.ville}"
        if dest:
            return text
        else:
            text += f"\nTél : {person.tel}"
            if person.mail == "":
                return text
            else :
                return text+f"\nMail : {person.mail}"        

