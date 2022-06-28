import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

class Letter:
    
    def __init__(self,root):
        self.root = root

    def save_letter(self, expediteur, destinataire, obj, msg):
        document = Document()
        sections = document.sections
        for section in sections:
            section.top_margin = Cm(3.81)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        p_expediteur = document.add_paragraph(self.check(expediteur))

        p_destinataire = document.add_paragraph(f"{self.check(destinataire,True)}\n\n")
        p_destinataire.paragraph_format.left_indent = Cm(9)
        # p_destinataire.paragraph_format.space_before = Cm(11)

        p_lieu = document.add_paragraph(
            f"{expediteur.ville.upper()},\nle {self.check_date()} ")
        p_lieu.paragraph_format.left_indent = Cm(9)

        if obj != "":
            p_objet = document.add_paragraph(f"Objet : {obj}")

        p_appel = document.add_paragraph(f"Madame, Monsieur,\n\n")
        msg_split = msg.split("\n")
        
        for text in msg_split:
            document.add_paragraph(text)

        p_signature = document.add_paragraph(
            f"{expediteur.prenom} {expediteur.nom}")
        p_signature.paragraph_format.left_indent = Cm(9)
        try:
            name = f"{datetime.datetime.now().strftime('%Y%m%d')}-{destinataire.nom}{destinataire.prenom}"
            extension = "docx"
            save_path = self.root.file_manager("save",name, extension)
            if not save_path == "":
                document.save(save_path)
            return True
        except:
            return False
    
    def check(self, person,dest=False):
        text = f"{person.nom} {person.prenom}\n{person.adresse}\n{person.cp} {person.ville.upper()}"
        if dest:
            return text
        else:
            text += f"\nTél : {person.tel}"
            if person.mail == "":
                return text
            else :
                return text+f"\nMail : {person.mail}"   
            
    def check_date(self):
        day = datetime.datetime.now().strftime('%d')
        year = datetime.datetime.now().strftime('%Y')
        month_int = int(datetime.datetime.now().strftime('%m'))
        month = ""
        match month_int:
            case 1:
                month = " Janvier "
            case 2:
                month = " Février "
            case 3:
                month = " Mars "
            case 4:
                month = " Avril "
            case 5:
                month = " Mai "
            case 6:
                month = " Juin "
            case 7:
                month = " Juillet "
            case 8:
                month = " Août "
            case 9:
                month = " Septembre "
            case 10:
                month = " Octobre "
            case 11:
                month = " Novembre "
            case 12:
                month = " Décembre "
        date = f"{day}{month}{year}"
        return date
        

